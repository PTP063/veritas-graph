"""
Veritas-Graph: Native OOXML Tracked Changes Engine with Schema Validation.
"""

import copy
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from lxml import etree

from src.exceptions import OOXMLInjectionError


class OOXMLRedliner:
    """Manipulates low-level OOXML AST to produce native Word revisions."""

    @staticmethod
    def _get_next_w_id(doc: DocumentType) -> int:
        """Inspects the entire document DOM to guarantee unique w:id generation.

        Args:
            doc: The python-docx Document object.

        Returns:
            The next available unique integer ID for Word track changes.
        """
        body = doc._body._element
        existing_ids = body.xpath("//@w:id")
        numeric_ids = [int(i) for i in existing_ids if i.isdigit()]
        return (max(numeric_ids) + 1) if numeric_ids else 1

    @classmethod
    def apply_tracked_change(
        cls,
        doc: DocumentType,
        paragraph_index: int,
        original_text: str,
        replacement_text: str,
        author: str = "Veritas-AI",
    ) -> None:
        """Injects native Word Tracked Changes AST nodes into the paragraph.

        Args:
            doc: The python-docx Document object.
            paragraph_index: The index of the target paragraph.
            original_text: The original text to be marked as deleted.
            replacement_text: The new text to be marked as inserted.
            author: Metadata author name for the track change.

        Raises:
            OOXMLInjectionError: If the target text is missing or spans multiple runs.
        """
        p = doc.paragraphs[paragraph_index]
        if original_text not in p.text:
            raise OOXMLInjectionError("Target text not found in paragraph.")

        w_id: int = cls._get_next_w_id(doc)
        now_iso: str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        # Build the exact Tracked Changes AST nodes
        del_node = etree.Element(
            qn("w:del"), nsmap={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        )
        del_node.set(qn("w:id"), str(w_id))
        del_node.set(qn("w:author"), author)
        del_node.set(qn("w:date"), now_iso)
        del_run = etree.SubElement(del_node, qn("w:r"))
        del_text = etree.SubElement(del_run, qn("w:delText"))
        del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        del_text.text = original_text

        ins_node = etree.Element(
            qn("w:ins"), nsmap={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        )
        ins_node.set(qn("w:id"), str(w_id + 1))
        ins_node.set(qn("w:author"), author)
        ins_node.set(qn("w:date"), now_iso)
        ins_run = etree.SubElement(ins_node, qn("w:r"))
        ins_text = etree.SubElement(ins_run, qn("w:t"))
        ins_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        ins_text.text = replacement_text

        for run in p.runs:
            if original_text in run.text:
                rPr = run._r.find(qn("w:rPr"))
                if rPr is not None:
                    # DEEPCOPY PREVENTS LXML DETACHMENT CORRUPTION
                    del_run.insert(0, copy.deepcopy(rPr))
                    ins_run.insert(0, copy.deepcopy(rPr))

                parts = run.text.split(original_text, 1)
                run.text = parts[0]

                run._r.addnext(del_node)
                del_node.addnext(ins_node)

                if parts[1]:
                    new_run_element = etree.Element(qn("w:r"))
                    if rPr is not None:
                        # DEEPCOPY FOR THE SUFFIX RUN
                        new_run_element.append(copy.deepcopy(rPr))
                    new_text = etree.SubElement(new_run_element, qn("w:t"))
                    new_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    new_text.text = parts[1]
                    ins_node.addnext(new_run_element)
                return

        # EXPLICIT FAILURE LOGGING FOR THE MVP GAP
        raise OOXMLInjectionError("Clause spans multiple <w:r> runs (Unsupported MVP edge case).")

    @classmethod
    def inject_warning_comment(
        cls, doc: DocumentType, paragraph_index: int, warning_text: str
    ) -> None:
        """Appends an in-document escalation marker for clauses requiring human review.

        Args:
            doc: The python-docx Document object.
            paragraph_index: The target paragraph index.
            warning_text: The text to inject as a warning.
        """
        p = doc.paragraphs[paragraph_index]
        run = p.add_run(f" [{warning_text}]")
        run.font.bold = True

    @classmethod
    def validate_document(cls, doc: DocumentType) -> bool:
        """Roundtrip verification: confirms generated XML parses cleanly.

        Args:
            doc: The python-docx Document object.

        Returns:
            True if the document can be successfully serialized and reloaded, else False.
        """
        try:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            reloaded = Document(buffer)
            return len(reloaded.paragraphs) > 0
        except Exception:
            return False
