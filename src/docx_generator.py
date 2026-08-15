from docx import Document

def generate_mock_nda(output_path: str):
    doc = Document()
    doc.add_heading('Non-Disclosure Agreement', 0)
    
    doc.add_paragraph(
        'This Non-Disclosure Agreement (the "Agreement") is entered into by and between the parties.'
    )
    
    doc.add_heading('1. Confidentiality', level=1)
    
    doc.add_paragraph(
        'The Receiving Party shall keep the Disclosing Party\'s information confidential for a period of two years.'
    )
    
    doc.save(output_path)
