"""
Veritas-Graph: Main Execution Entry Point
Runs the end-to-end pipeline and outputs system telemetry.
"""
import asyncio
import os
import sys
from docx import Document
from google import genai
from dotenv import load_dotenv  # <-- ADD THIS

# Load the variables from the .env file into the OS environment
load_dotenv()                   # <-- ADD THIS

from src.graph import VeritasGraphOrchestrator
from src.pipeline import VeritasPipeline

async def main():
    print("\n" + "="*60)
# ... [rest of your code stays exactly the same] ...

    print("🚀 INITIALIZING VERITAS-GRAPH PIPELINE 🚀")
    print("="*60 + "\n")

    # 1. API Key Check
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("❌ ERROR: GEMINI_API_KEY environment variable not set.")
        print("Run: export GEMINI_API_KEY='your_key_here' (Mac/Linux)")
        print("Or:  set GEMINI_API_KEY=your_key_here (Windows)")
        sys.exit(1)

    # 2. Prepare Demo Data (Creating a mock contract on the fly)
    print("📄 Preparing mock corporate contract (mock_contract.docx)...")
    doc = Document()
    doc.add_heading("Master Services Agreement", 0)
    doc.add_paragraph("This agreement governs the terms of service.")
    
    # We include mixed formatting here to prove the deepcopy AST injection works
    p = doc.add_paragraph("Section 1. Change of Control. ")
    target_run = p.add_run("In the event the company is acquired, all active obligations are immediately null and void, and a $5,000,000 penalty shall be assessed.")
    target_run.bold = True
    doc.save("mock_contract.docx")

    # The raw text extraction that Phase 1 (pdfplumber) would normally provide
    raw_text = (
        "Master Services Agreement\n"
        "This agreement governs the terms of service.\n"
        "Section 1. Change of Control. In the event the company is acquired, all active obligations "
        "are immediately null and void, and a $5,000,000 penalty shall be assessed."
    )

    # 3. Initialize Pipeline
    print("🤖 Booting Gemini 3.1 Pro Orchestrator...")
    client = genai.Client(api_key=api_key)
    # Capping concurrency to 3 to respect standard API rate limits
    orchestrator = VeritasGraphOrchestrator(client, max_concurrency=3)
    pipeline = VeritasPipeline(orchestrator)

    # 4. Execute the Graph
    print("⚡ Executing Asynchronous DAG Fan-Out...")
    final_state = await pipeline.process_contract(
        document_text=raw_text, 
        docx_path="mock_contract.docx"
    )

    # 5. Output Telemetry & Results
    print("\n" + "="*60)
    print("✅ PIPELINE EXECUTION COMPLETE ✅")
    print("="*60)
    
    print("\n📊 TELEMETRY:")
    print(f"  - Node Latency: {final_state.node_latency_ms.get('full_dag_fanout', 0):.2f} ms")
    print(f"  - Total Tokens (Estimated): {final_state.token_telemetry.total_tokens}")
    
    print("\n🔍 AUDIT LOG:")
    if final_state.failed_chunks:
        print(f"  - ⚠️ Failed Injections: {len(final_state.failed_chunks)}")
        for f in final_state.failed_chunks:
            print(f"    * {f.get('error')}")
    else:
        print("  - 🟢 Failed Injections: 0")

    print(f"  - 🟢 Successful Redlines: {len(final_state.redlines)}")
    for r in final_state.redlines:
        print(f"    * Replaced: '{r.original_text[:35]}...' -> '{r.replacement_text[:35]}...'")

    if final_state.requires_global_human_review:
        print("\n🚨 SYSTEM WARNING: Global Human Review Flagged!")

    print(f"\n📂 OUTPUT: Native Tracked Changes saved to: redlined_{final_state.document_id}.docx")
    print("Open this file in Microsoft Word to see the native AST injection in action!\n")

if __name__ == "__main__":
    asyncio.run(main())