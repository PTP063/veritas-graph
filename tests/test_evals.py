"""
Veritas-Graph: Adversarial Evaluation & Benchmarking Suite.
"""
import pytest
from docx import Document
from src.chunker import LegalChunker
from src.aligner import DocumentAligner
from src.redliner import OOXMLRedliner


def test_unstructured_document_flag():
    chunker = LegalChunker()
    # Feed arbitrary text with no legal headers
    garbage_text = "This is a random paragraph without any section headers." * 200
    chunks, is_unstructured = chunker.chunk(garbage_text)
    assert is_unstructured is True
    assert len(chunks) > 1


def test_aligner_disambiguation():
    doc = Document()
    doc.add_paragraph("Section 1. Confidentiality obligations shall last two years.")
    doc.add_paragraph("Section 2. Confidentiality obligations shall last two years.")

    # Disambiguate identical strings via preceding context
    idx, score, is_ambiguous = DocumentAligner.find_best_paragraph_match(
        doc, 
        target_text="Confidentiality obligations shall last two years.",
        context_before="Section 2."
    )
    assert idx == 1


def test_ooxml_validation_roundtrip():
    doc = Document()
    doc.add_paragraph("Target paragraph to be changed.")
    
    OOXMLRedliner.apply_tracked_change(
        doc, 
        paragraph_index=0, 
        original_text="Target paragraph to be changed.", 
        replacement_text="Revised paragraph."
    )
    assert OOXMLRedliner.validate_document(doc) is True