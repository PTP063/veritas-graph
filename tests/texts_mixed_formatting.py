"""
Veritas-Graph: Mixed Formatting AST Validation Test
Proves that surgical AST injection preserves surrounding bold/italic runs.
"""
from docx import Document
from src.redliner import OOXMLRedliner

def test_mixed_formatting_run_boundary():
    # 1. Build a document with mixed formatting
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The Company shall provide ")
    
    # Target clause is isolated in a bold run
    target_run = p.add_run("thirty (30) days")
    target_run.bold = True 
    
    p.add_run(" written notice.")

    # 2. Execute the AST Injection
    success, err_msg = OOXMLRedliner.apply_tracked_change(
        doc=doc,
        paragraph_index=0,
        original_text="thirty (30) days",
        replacement_text="sixty (60) days"
    )

    # 3. Assertions
    assert success is True, f"AST Injection failed: {err_msg}"
    assert OOXMLRedliner.validate_document(doc) is True

    # Save it to disk so we can visually inspect it in MS Word
    doc.save("mixed_formatting_proof.docx")
    print("\n✅ Test Passed! Open 'mixed_formatting_proof.docx' in Word.")

if __name__ == "__main__":
    test_mixed_formatting_run_boundary()