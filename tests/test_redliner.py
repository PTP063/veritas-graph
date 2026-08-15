import os
import tempfile
import pytest
from src.docx_generator import generate_mock_nda
from src.redliner import OOXMLRedliner
from docx import Document

def test_ast_redlining():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
        in_path = tmp_in.name
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_out:
        out_path = tmp_out.name
        
    try:
        generate_mock_nda(in_path)
        
        target = "two years"
        revised = "five years"
        
        doc = Document(in_path)
        
        for i, p in enumerate(doc.paragraphs):
            if target in p.text:
                OOXMLRedliner.apply_tracked_change(doc, i, target, revised)
                break
                
        doc.save(out_path)
        
        # Reload the output and verify AST
        doc = Document(out_path)
        
        found_del = False
        found_ins = False
        
        for p in doc.paragraphs:
            xml = p._p.xml
            if '<w:del ' in xml and 'two years' in xml:
                found_del = True
                assert 'w:author="Veritas-AI"' in xml
            if '<w:ins ' in xml and 'five years' in xml:
                found_ins = True
                assert 'w:author="Veritas-AI"' in xml
                
        assert found_del, "w:del node not found in AST"
        assert found_ins, "w:ins node not found in AST"
        
    finally:
        if os.path.exists(in_path):
            os.remove(in_path)
        if os.path.exists(out_path):
            os.remove(out_path)
